"""Word extractor.

Timesheets arrive in Word as either a table (most common) or prose. We parse every
table through the shared grid parser and keep the one that yields the most rows; if no
table is usable we fall back to the paragraph text through the email/prose parser.

  .docx → python-docx (tables + paragraphs)
  .doc  → legacy OLE binary: best-effort readable-text extraction via olefile, then
          the prose parser. The old .doc format interleaves text with control
          structures, so this is lossy — but it recovers EMP IDs / "N days" patterns,
          and an unreadable doc degrades to escalate rather than crashing.
"""

from __future__ import annotations

import re
from pathlib import Path

from ..schema import TimesheetExtraction
from . import email as email_ex
from .excel import parse_grid


def extract_docx(path: str | Path) -> TimesheetExtraction:
    import io

    import docx

    d = docx.Document(io.BytesIO(Path(path).read_bytes()))

    best = TimesheetExtraction()
    for table in d.tables:
        grid = [[cell.text for cell in row.cells] for row in table.rows]
        r = parse_grid(grid)
        if len(r.rows) > len(best.rows):
            best = r

    prose = "\n".join(p.text for p in d.paragraphs if p.text and p.text.strip())

    if best.rows:
        # supplement client / period from the surrounding prose when the table omitted them
        if (not best.client_hint or not best.period) and prose.strip():
            sup = email_ex.extract_email(prose)
            best.client_hint = best.client_hint or sup.client_hint
            best.period = best.period or sup.period
        return best

    # no usable table → treat the document text like an email/structured body
    return email_ex.extract_email(prose) if prose.strip() else TimesheetExtraction()


def _doc_readable_text(path: str | Path) -> str:
    """Best-effort text from a legacy .doc WordDocument stream (lossy)."""
    try:
        import olefile

        if not olefile.isOleFile(str(path)):
            return ""
        with olefile.OleFileIO(str(path)) as ole:
            if not ole.exists("WordDocument"):
                return ""
            raw = ole.openstream("WordDocument").read()
    except Exception:  # noqa: BLE001
        return ""
    # pull runs of printable latin-1 and utf-16le text out of the binary stream
    latin = re.sub(rb"[^\x09\x0a\x0d\x20-\x7e]+", b" ", raw).decode("latin-1", "ignore")
    try:
        utf16 = raw.decode("utf-16-le", "ignore")
        utf16 = re.sub(r"[^\x09\x0a\x0d\x20-\uffff]+", " ", utf16)
    except Exception:  # noqa: BLE001
        utf16 = ""
    text = utf16 if len(utf16) > len(latin) else latin
    return re.sub(r"\s{2,}", " ", text).strip()


def extract_doc(path: str | Path) -> TimesheetExtraction:
    text = _doc_readable_text(path)
    return email_ex.extract_email(text) if text else TimesheetExtraction()

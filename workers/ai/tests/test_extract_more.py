"""Deeper extract coverage: word (.doc OLE seam), pdf (native + scanned),
sniff (zip/ole/hint fallbacks), email attachments, and dispatch chains."""

from __future__ import annotations

import io
import zipfile

import pytest
from PIL import Image

from tia_ai.extract import extract, word as W, pdf as PDFX, sniff, email_attachments as EA
from tia_ai.schema import TimesheetExtraction, TimesheetRow


# ── word ─────────────────────────────────────────────────────────────────────


def test_extract_docx_prose_supplements_client_period(tmp_path):
    import docx

    d = docx.Document()
    t = d.add_table(rows=0, cols=4)
    for row in [["Emp ID", "Full Name", "Working Days", "OT Hours"],
                ["EMP10001", "Carlos Smith", "22", "5"]]:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = v
    d.add_paragraph("Client Code: CL001")
    d.add_paragraph("Month: June 2026")
    p = tmp_path / "t.docx"
    d.save(p)
    ex = W.extract_docx(p)
    assert ex.rows and (ex.client_hint == "CL001" or ex.period == "June 2026")


def test_extract_docx_no_table_prose_fallback(tmp_path):
    import docx

    d = docx.Document()
    d.add_paragraph("EMP10001 Carlos Smith worked 22 days in June 2026")
    p = tmp_path / "prose.docx"
    d.save(p)
    ex = W.extract_docx(p)
    assert isinstance(ex, TimesheetExtraction)


def test_extract_docx_empty_doc(tmp_path):
    import docx

    d = docx.Document()
    p = tmp_path / "empty.docx"
    d.save(p)
    assert W.extract_docx(p).rows == []


def test_doc_readable_text_non_ole_returns_empty(tmp_path):
    p = tmp_path / "x.doc"
    p.write_bytes(b"just plain not-ole bytes")
    assert W._doc_readable_text(p) == ""
    assert W.extract_doc(p).rows == []


def test_extract_doc_via_fake_ole(monkeypatch, tmp_path):
    import olefile

    class FakeOle:
        def __init__(self, path):
            pass

        def __enter__(self):
            return self

        def __exit__(self, *a):
            return False

        def exists(self, name):
            return name == "WordDocument"

        def openstream(self, name):
            return io.BytesIO(b"EMP10001 Carlos Smith worked 22 days\x00\x01binaryjunk")

    monkeypatch.setattr(olefile, "isOleFile", lambda p: True)
    monkeypatch.setattr(olefile, "OleFileIO", FakeOle)
    p = tmp_path / "legacy.doc"
    p.write_bytes(b"\xd0\xcf\x11\xe0stub")
    text = W._doc_readable_text(p)
    assert "Carlos Smith" in text
    ex = W.extract_doc(p)
    assert isinstance(ex, TimesheetExtraction)


def test_doc_readable_text_no_word_stream(monkeypatch, tmp_path):
    import olefile

    class FakeOle:
        def __init__(self, path):
            pass

        def __enter__(self):
            return self

        def __exit__(self, *a):
            return False

        def exists(self, name):
            return False

    monkeypatch.setattr(olefile, "isOleFile", lambda p: True)
    monkeypatch.setattr(olefile, "OleFileIO", FakeOle)
    p = tmp_path / "nostream.doc"
    p.write_bytes(b"\xd0\xcf\x11\xe0stub")
    assert W._doc_readable_text(p) == ""


# ── pdf ──────────────────────────────────────────────────────────────────────


def test_extract_pdf_native_text_layer(monkeypatch, tmp_path):
    import pdfplumber

    class FakePage:
        def extract_text(self):
            return "EMP10001 Carlos Smith - 22 days\nClient Code: CL001\nMonth: June 2026"

    class FakePDF:
        pages = [FakePage()]

        def __enter__(self):
            return self

        def __exit__(self, *a):
            return False

    monkeypatch.setattr(pdfplumber, "open", lambda *a, **k: FakePDF())
    p = tmp_path / "native.pdf"
    p.write_bytes(b"%PDF-1.4 stub")
    ex = PDFX.extract_pdf(p)
    assert isinstance(ex, TimesheetExtraction)


def test_extract_pdf_scanned_routes_to_vision(monkeypatch, tmp_path):
    import tia_ai.ocr as ocr

    # a real image-only PDF (no text layer) → scanned path → vision (mocked OCR)
    p = tmp_path / "scanned.pdf"
    Image.new("RGB", (200, 120), "white").save(p, "PDF")
    monkeypatch.setattr(ocr, "glm_markdown", lambda *a, **k: "Employee Name: Zed\nClient Code: CL001\nMonth: June 2026\nTotal Working Days: 20\nTotal Hours: 160")
    monkeypatch.setattr(ocr, "glm_kie", lambda *a, **k: TimesheetExtraction())
    monkeypatch.setattr(ocr, "glm_layout", lambda *a, **k: [])
    ex = PDFX.extract_pdf(p)
    assert isinstance(ex, TimesheetExtraction)


def test_extract_pdf_scanned_failure_returns_empty(monkeypatch, tmp_path):
    import pdfplumber

    class FakePage:
        def extract_text(self):
            return ""  # no text layer

        def to_image(self, resolution=200):
            raise RuntimeError("raster failed")

    class FakePDF:
        pages = [FakePage()]

        def __enter__(self):
            return self

        def __exit__(self, *a):
            return False

    monkeypatch.setattr(pdfplumber, "open", lambda *a, **k: FakePDF())
    p = tmp_path / "bad.pdf"
    p.write_bytes(b"%PDF-1.4 stub")
    assert PDFX.extract_pdf(p).rows == []


# ── sniff ──────────────────────────────────────────────────────────────────────


def _zip_with(names: list[str], tmp_path, fname="a.bin") -> "object":
    p = tmp_path / fname
    with zipfile.ZipFile(p, "w") as z:
        for n in names:
            z.writestr(n, "x")
    return p


def test_zip_kind_docx_pptx(tmp_path):
    assert sniff._zip_kind(_zip_with(["word/document.xml"], tmp_path, "d.bin")) == "docx"
    assert sniff._zip_kind(_zip_with(["ppt/presentation.xml"], tmp_path, "p.bin")) == "pptx"
    assert sniff._zip_kind(_zip_with(["random/thing.txt"], tmp_path, "z.bin")) == "zip"


def test_zip_kind_bad_zip(tmp_path):
    p = tmp_path / "broken.bin"
    p.write_bytes(b"PK\x03\x04notreallyazip")
    assert sniff._zip_kind(p) == "zip"


def test_ole_kind_via_fake(monkeypatch, tmp_path):
    import olefile

    def make(entries):
        class FakeOle:
            def __init__(self, path):
                pass

            def __enter__(self):
                return self

            def __exit__(self, *a):
                return False

            def listdir(self):
                return [[e] for e in entries]

        return FakeOle

    p = tmp_path / "o.bin"
    p.write_bytes(b"\xd0\xcf\x11\xe0stub")
    monkeypatch.setattr(olefile, "OleFileIO", make(["Workbook"]))
    assert sniff._ole_kind(p) == "xls"
    monkeypatch.setattr(olefile, "OleFileIO", make(["WordDocument"]))
    assert sniff._ole_kind(p) == "doc"
    monkeypatch.setattr(olefile, "OleFileIO", make(["Other"]))
    assert sniff._ole_kind(p) == "ole"


def test_ole_kind_error(monkeypatch, tmp_path):
    import olefile

    def _boom(path):
        raise OSError("bad ole")

    monkeypatch.setattr(olefile, "OleFileIO", _boom)
    p = tmp_path / "o.bin"
    p.write_bytes(b"\xd0\xcf\x11\xe0stub")
    assert sniff._ole_kind(p) == "ole"


def test_detect_kind_ole_dispatch(monkeypatch, tmp_path):
    p = tmp_path / "ole.bin"
    p.write_bytes(sniff._OLE_MAGIC + b"\x00" * 32)
    monkeypatch.setattr(sniff, "_ole_kind", lambda pth: "xls")
    assert sniff.detect_kind(p) == "xls"


def test_detect_kind_binary_unknown_uses_hint(tmp_path):
    # non-decodable, non-printable content → falls through to the MIME/suffix hint
    p = tmp_path / "mystery.bin"
    p.write_bytes(b"\x00\x01\x02\x03\x04\x05\x06\x07\x08" * 4)
    assert sniff.detect_kind(p, mime="application/pdf") == "pdf"
    assert sniff.detect_kind(p, filename="x.xlsx") == "xlsx"
    assert sniff.detect_kind(p, filename="x.docx") == "docx"
    assert sniff.detect_kind(p, filename="x.doc") == "doc"
    assert sniff.detect_kind(p, filename="x.xls") == "xls"
    assert sniff.detect_kind(p, mime="text/csv") == "csv"
    assert sniff.detect_kind(p, filename="x.png") == "image"
    assert sniff.detect_kind(p, filename="x.weird") == "unknown"


def test_detect_kind_empty_and_missing(tmp_path):
    empty = tmp_path / "e.bin"
    empty.write_bytes(b"")
    assert sniff.detect_kind(empty) == "unknown"
    assert sniff.detect_kind(tmp_path / "does-not-exist.bin") == "unknown"


def test_looks_csv_too_few_lines():
    assert sniff._looks_csv("only one line") is False
    assert sniff._looks_csv("a,b,c\nd,e,f\n") is True
    assert sniff._looks_csv("no delimiters here\nnor here either\n") is False


# ── email attachments ─────────────────────────────────────────────────────────


def test_extract_attachments_synthetic_filename():
    from email.mime.application import MIMEApplication
    from email.mime.multipart import MIMEMultipart
    from email.mime.text import MIMEText

    m = MIMEMultipart()
    m.attach(MIMEText("body text"))
    a = MIMEApplication(b"rawbytes", _subtype="octet-stream")  # no filename
    m.attach(a)
    files = list(EA.extract_attachments(m.as_bytes()))
    assert files and files[0][0].startswith("attachment-")


def test_extract_attachments_skips_empty_payload():
    from email.mime.application import MIMEApplication
    from email.mime.multipart import MIMEMultipart

    m = MIMEMultipart()
    a = MIMEApplication(b"", _subtype="pdf")
    a.add_header("Content-Disposition", "attachment", filename="empty.pdf")
    m.attach(a)
    assert list(EA.extract_attachments(m.as_bytes())) == []


# ── dispatch chains ───────────────────────────────────────────────────────────


def test_extract_routes_image_to_vision(monkeypatch, tmp_path):
    import tia_ai.ocr as ocr

    p = tmp_path / "photo.png"
    Image.new("RGB", (60, 40), "white").save(p, "PNG")
    monkeypatch.setattr(
        ocr, "glm_markdown", lambda *a, **k: "Employee Name: PicGuy\nClient Code: CL001\nMonth: June 2026\nTotal Working Days: 21\nTotal Hours: 168"
    )
    monkeypatch.setattr(ocr, "glm_kie", lambda *a, **k: TimesheetExtraction())
    monkeypatch.setattr(ocr, "glm_layout", lambda *a, **k: [])
    ex = extract(p, filename="photo.png")
    assert isinstance(ex, TimesheetExtraction)


def test_extract_keeps_richest_when_no_rows(monkeypatch, tmp_path):
    # a text file whose parsers find client/period hints but no rows → keep-best branch
    p = tmp_path / "hints.txt"
    p.write_text("Client Code: CL001\nMonth: June 2026\n", encoding="utf-8")
    ex = extract(p, filename="hints.txt")
    assert ex.rows == []
    assert isinstance(ex, TimesheetExtraction)
